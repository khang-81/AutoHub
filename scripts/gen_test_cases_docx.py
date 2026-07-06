"""
Script tạo file DOCX bảng kiểm thử RÚT GỌN cho dự án AutoHub Rent-A-Car.
Chỉ giữ các test case chính (~60 cases thay vì 114).
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"C:\Do_an\Rent-A-Car\docs\Bang_Kiem_Thu_AutoHub.docx"
OUTPUT_TMP = r"C:\Do_an\Rent-A-Car\docs\Bang_Kiem_Thu_AutoHub_new.docx"

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for border_name in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tc_borders.append(border)
    tc_pr.append(tc_borders)

def add_test_table(doc, headers, rows, header_color="2E5090"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    col_widths = [Cm(2.0), Cm(5.0), Cm(4.5), Cm(5.5), Cm(1.8)]
    for i, w in enumerate(col_widths):
        for cell in table.columns[i].cells:
            cell.width = w
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        set_cell_bg(cell, header_color)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(11)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    return table

def add_module_heading(doc, text, module_id):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"{module_id}. {text}")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x50, 0x90)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

sections = doc.sections
for s in sections:
    s.top_margin = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin = Cm(2)
    s.right_margin = Cm(2)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("BẢNG KIỂM THỬ CHỨC NĂNG")
run.font.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Dự án: AutoHub – Hệ thống cho thuê & mua bán xe ô tô")
run.font.bold = True
run.font.size = Pt(13)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("Đồ án tốt nghiệp (Phiên bản rút gọn - các chức năng chính)")
run.font.italic = True
run.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Công nghệ sử dụng:")
run.font.bold = True
doc.add_paragraph("• Frontend: React 19, Vite, TypeScript, Tailwind CSS, Zustand, React Query", style='List Bullet')
doc.add_paragraph("• Backend: Spring Boot 3, Spring Security, JWT, JPA/Hibernate, Flyway", style='List Bullet')
doc.add_paragraph("• Database: SQL Server 2022", style='List Bullet')

p = doc.add_paragraph()
run = p.add_run("Ký hiệu kết quả:")
run.font.bold = True
doc.add_paragraph("• Pass: Đạt  |  • Fail: Không đạt  |  • N/A: Không áp dụng", style='List Bullet')

doc.add_paragraph()
note = doc.add_paragraph()
run = note.add_run("Tổng số test case: ~60  |  Số module: 20")
run.font.bold = True
run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

doc.add_page_break()

HEADERS = ["ID", "Mô tả", "Dữ liệu đầu vào", "Kỳ vọng", "Kết quả"]

# TC01: AUTH (5)
add_module_heading(doc, "Xác thực & Phân quyền (Authentication)", "TC01")
add_test_table(doc, HEADERS, [
    ["TC01_01", "Đăng ký tài khoản mới",
     "Email: newuser@test.com | Password: Test1234@",
     "Tạo tài khoản thành công, gán role 'user', chuyển về trang Login",
     "Pass"],
    ["TC01_02", "Đăng nhập thành công",
     "Email: user@autohub.id.vn | Password: admin123@",
     "Cấp JWT token, lưu localStorage, redirect về trang chủ",
     "Pass"],
    ["TC01_03", "Đăng nhập sai mật khẩu",
     "Email: user@autohub.id.vn | Password: wrongPassword",
     "Hiển thị 'Email hoặc mật khẩu không chính xác'",
     "Pass"],
    ["TC01_04", "Quên mật khẩu + Reset bằng OTP",
     "Email user có thật → nhận OTP 6 số → nhập OTP + mật khẩu mới",
     "Cập nhật MK mới, tăng tokenVersion (vô hiệu JWT cũ), redirect Login",
     "Pass"],
    ["TC01_05", "Đổi mật khẩu trong dashboard",
     "MK cũ: admin123@ → MK mới: NewPass123@",
     "Cập nhật thành công, các phiên khác bị logout, hiện toast thành công",
     "Pass"],
])

# TC02: KYC (3)
add_module_heading(doc, "Xác minh danh tính (KYC)", "TC02")
add_test_table(doc, HEADERS, [
    ["TC02_01", "Upload CCCD + GPLX thành công",
     "Trang /dashboard/kyc | File: cccd.jpg, gplx.jpg (PNG/JPG, <5MB)",
     "Upload thành công, status tự chuyển APPROVED (dev auto-approve)",
     "Pass"],
    ["TC02_02", "Admin duyệt/từ chối KYC",
     "Admin /admin/viewing-appointments | Chọn user | Duyệt hoặc Từ chối + nhập lý do",
     "Cập nhật status, gửi email thông báo cho user, lưu adminNote",
     "Pass"],
    ["TC02_03", "Chặn thuê xe khi chưa KYC",
     "User status: NOT_SUBMITTED | Click 'Thuê xe'",
     "Chặn + thông báo 'Cần xác minh danh tính trước', redirect /dashboard/kyc",
     "Pass"],
])

# TC03: CAR CATALOG (4)
add_module_heading(doc, "Danh mục xe (Car Catalog)", "TC03")
add_test_table(doc, HEADERS, [
    ["TC03_01", "Xem danh sách xe cho thuê",
     "Truy cập /cars",
     "Hiển thị grid xe với ảnh, tên, giá/ngày, hãng; phân trang 12 xe/trang",
     "Pass"],
    ["TC03_02", "Tìm kiếm xe theo từ khóa",
     "Search: 'Toyota'",
     "Hiển thị xe có tên/brand chứa 'Toyota' (case-insensitive)",
     "Pass"],
    ["TC03_03", "Lọc xe theo hãng + giá",
     "Brand: 'Mercedes' | minPrice: 500K | maxPrice: 2M",
     "Chỉ hiển thị xe Mercedes có giá trong khoảng [500K, 2M]",
     "Pass"],
    ["TC03_04", "Xem chi tiết xe",
     "Click vào 1 xe",
     "Mở /cars/{id}: gallery, thông số, giá, 2 nút 'Thuê' & 'Mua'",
     "Pass"],
])

# TC04: VIEWING APPOINTMENTS (4) - theo mẫu
add_module_heading(doc, "Đặt lịch xem xe (Viewing Appointments)", "TC04")
add_test_table(doc, HEADERS, [
    ["TC04_01", "Hiển thị danh sách sự kiện",
     "Người dùng đã tạo ít nhất 1 sự kiện",
     "Hệ thống hiển thị danh sách sự kiện theo 4 tab: Sắp diễn ra, Đã qua, Bị từ chối, Chờ duyệt",
     "Pass"],
    ["TC04_02", "Người dùng chưa có sự kiện nào",
     "Người dùng chưa từng tạo sự kiện",
     "Hệ thống hiển thị thông báo 'Bạn chưa tạo sự kiện nào'",
     "Pass"],
    ["TC04_03", "Lọc sự kiện theo từ khóa",
     "Nhập từ khóa vào ô tìm kiếm (ví dụ: 'Mercedes')",
     "Hiển thị danh sách sự kiện phù hợp với từ khóa",
     "Pass"],
    ["TC04_04", "Chuyển tab sự kiện",
     "Chọn tab 'Đã qua'",
     "Hệ thống phản hồi trong vòng 2s và hiển thị đúng danh sách",
     "Pass"],
])

# TC05: RENTAL (5)
add_module_heading(doc, "Thuê xe (Rental)", "TC05")
add_test_table(doc, HEADERS, [
    ["TC05_01", "Chọn ngày thuê hợp lệ + bảo hiểm",
     "Pickup: 2026-07-01 | Return: 2026-07-05 | Gói STANDARD",
     "Tổng = giá/ngày × số ngày + phí BH, calendar disable ngày đã thuê",
     "Pass"],
    ["TC05_02", "Chọn ngày không hợp lệ",
     "Pickup: 2026-07-05 | Return: 2026-07-01",
     "Báo lỗi 'Ngày trả phải sau ngày nhận', không cho submit",
     "Pass"],
    ["TC05_03", "Áp dụng mã khuyến mãi hợp lệ",
     "Nhập mã 'SUMMER2026' (active)",
     "Áp dụng discount 15% vào subtotal, usedCount++",
     "Pass"],
    ["TC05_04", "Tạo đơn thuê thành công",
     "User đã KYC | Đã chọn ngày + BH | Click 'Đặt xe'",
     "Tạo Rental PENDING_PAYMENT + Invoice tự động, redirect payment page",
     "Pass"],
    ["TC05_05", "Hủy đơn thuê",
     "Đơn PENDING_PAYMENT | Click 'Hủy' + nhập lý do",
     "Áp dụng cancel policy, hoàn cọc theo depositRefundRatio, status CANCELLED",
     "Pass"],
])

# TC06: PAYMENT (3)
add_module_heading(doc, "Thanh toán thuê xe (Bank Transfer)", "TC06")
add_test_table(doc, HEADERS, [
    ["TC06_01", "Hiển thị QR + thông tin CK",
     "Trang /dashboard/payment/{id} (PENDING_PAYMENT)",
     "Hiển thị QR VietQR, STK, NH, chủ TK, số tiền, nội dung CK",
     "Pass"],
    ["TC06_02", "User submit đã chuyển khoản",
     "Click 'Tôi đã chuyển khoản'",
     "Status → PENDING_CONFIRM, gửi email admin",
     "Pass"],
    ["TC06_03", "Admin xác nhận thanh toán",
     "Admin /admin/cars/rent?tab=orders | Click 'Xác nhận'",
     "Status → CONFIRMED, gửi email user",
     "Pass"],
])

# TC07: SALE (3)
add_module_heading(doc, "Mua xe (Sale Orders)", "TC07")
add_test_table(doc, HEADERS, [
    ["TC07_01", "Mua xe thành công",
     "Trang /cars/{id}/buy | Xe AVAILABLE | Click 'Mua ngay'",
     "Tạo SaleOrder PENDING_PAYMENT, redirect payment page",
     "Pass"],
    ["TC07_02", "Race condition: 2 user cùng mua",
     "User A và B click 'Mua' cùng lúc",
     "1 user thành công, user còn lại lỗi 'Xe đã bán' (@Version conflict)",
     "Pass"],
    ["TC07_03", "Hủy đơn mua",
     "SaleOrder PENDING_PAYMENT | Click 'Hủy'",
     "Status CANCELLED, xe AVAILABLE trở lại",
     "Pass"],
])

# TC08: REVIEWS (3)
add_module_heading(doc, "Đánh giá (Reviews)", "TC08")
add_test_table(doc, HEADERS, [
    ["TC08_01", "Tạo đánh giá sau khi thuê",
     "User đã hoàn thành rental | Rating: 5 | Comment: 'Xe rất tốt'",
     "Review tạo thành công, hiển thị trên trang chi tiết xe",
     "Pass"],
    ["TC08_02", "Xem đánh giá của xe + filter rating",
     "Trang /cars/{id} → section 'Đánh giá' | Filter minRating=4",
     "Hiển thị rating trung bình + danh sách review ≥4 sao",
     "Pass"],
    ["TC08_03", "Admin ẩn/reply đánh giá",
     "Admin /admin/reviews | Click 'Ẩn' hoặc 'Reply'",
     "Review bị ẩn hoặc reply hiển thị trên frontend",
     "Pass"],
])

# TC09: PROMOTIONS (2)
add_module_heading(doc, "Khuyến mãi (Promotions)", "TC09")
add_test_table(doc, HEADERS, [
    ["TC09_01", "Admin tạo mã khuyến mãi",
     "Admin /admin/promotions | Mã: 'SUMMER2026' | Discount: 15% | Có hiệu lực",
     "Mã được tạo, có thể áp dụng trong flow thuê/mua",
     "Pass"],
    ["TC09_02", "Mã hết hạn không áp dụng được",
     "Mã 'OLDCODE' đã quá endDate",
     "Báo lỗi 'Mã khuyến mãi đã hết hạn'",
     "Pass"],
])

# TC10: PROFILE (2)
add_module_heading(doc, "Hồ sơ cá nhân (Profile)", "TC10")
add_test_table(doc, HEADERS, [
    ["TC10_01", "Cập nhật thông tin cá nhân",
     "Sửa Họ tên + SĐT + Địa chỉ | Click 'Lưu'",
     "Cập nhật thành công, hiển thị toast, thông tin mới hiển thị ngay",
     "Pass"],
    ["TC10_02", "Upload avatar",
     "Chọn avatar.jpg | Click 'Upload'",
     "Avatar cập nhật, hiển thị trong Navbar và profile",
     "Pass"],
])

# TC11: CONTACT (2)
add_module_heading(doc, "Liên hệ & Trang tĩnh", "TC11")
add_test_table(doc, HEADERS, [
    ["TC11_01", "Gửi form liên hệ",
     "Trang /contact | Họ tên + Email + SĐT + Nội dung",
     "Gửi email admin thành công, hiển thị toast, reset form",
     "Pass"],
    ["TC11_02", "Rate limit form liên hệ",
     "Gửi 6 lần trong 1 phút (limit 5/phút)",
     "Lần 6 bị chặn với HTTP 429",
     "Pass"],
])

# TC12: AI CHATBOT (3)
add_module_heading(doc, "AI Chatbot (Arcanic)", "TC12")
add_test_table(doc, HEADERS, [
    ["TC12_01", "Mở AI Chatbot và hỏi về xe",
     "Click nút AI góc dưới phải | Hỏi: 'Xe 5 chỗ giá dưới 1 triệu?'",
     "Chatbot mở, gọi Arcanic API với system prompt có context xe, trả lời gợi ý",
     "Pass"],
    ["TC12_02", "AI ẩn trên admin",
     "Truy cập /admin/*",
     "AI chatbot ẩn hoàn toàn (isAdminRoute check)",
     "Pass"],
    ["TC12_03", "Kiểm tra trạng thái AI",
     "GET /api/ai/status",
     "Trả về { aiConfigured } tùy ARCANIC_API_KEY có hay không",
     "Pass"],
])

# TC13: DASHBOARD & REPORTS (4)
add_module_heading(doc, "Admin Dashboard & Báo cáo", "TC13")
add_test_table(doc, HEADERS, [
    ["TC13_01", "Xem thống kê tổng quan",
     "Admin /admin",
     "Hiển thị 4 stat cards thuê + 3 stat cards bán + đơn gần đây + phân bố hãng",
     "Pass"],
    ["TC13_02", "Xem biểu đồ doanh thu",
     "Trang /admin/reports",
     "Bar chart doanh thu thuê + bán theo tháng (Recharts)",
     "Pass"],
    ["TC13_03", "Export báo cáo Excel",
     "Click 'Export Excel'",
     "Tải file .xlsx danh sách rentals theo tháng (Apache POI)",
     "Pass"],
    ["TC13_04", "Export báo cáo PDF",
     "Click 'Export PDF'",
     "Tải file .pdf với header logo + tiêu đề (OpenPDF)",
     "Pass"],
])

# TC14: MANAGE CARS (3)
add_module_heading(doc, "Admin – Quản lý xe (Manage Cars)", "TC14")
add_test_table(doc, HEADERS, [
    ["TC14_01", "Thêm xe cho thuê mới",
     "Admin /admin/cars/rent | Click 'Thêm xe' | Điền form đầy đủ",
     "Xe mới tạo với status AVAILABLE, listing RENT, hiển thị danh sách",
     "Pass"],
    ["TC14_02", "Sửa thông tin + upload ảnh",
     "Click 'Sửa' | Đổi giá | Upload 5 ảnh, sắp xếp thứ tự",
     "Cập nhật thành công, ảnh mới hiển thị trên trang public",
     "Pass"],
    ["TC14_03", "Xóa xe đang có đơn thuê",
     "Click 'Xóa' trên xe có rental CONFIRMED",
     "Chặn xóa với lỗi 'Không thể xóa xe đang có đơn thuê'",
     "Pass"],
])

# TC15: MANAGE BRANDS (2)
add_module_heading(doc, "Admin – Quản lý hãng xe", "TC15")
add_test_table(doc, HEADERS, [
    ["TC15_01", "Thêm hãng xe mới",
     "/admin/brands | Tên: 'VinFast' | Upload logo",
     "Brand tạo thành công, hiển thị trong dropdown filter xe",
     "Pass"],
    ["TC15_02", "Xóa hãng đang có xe",
     "Click 'Xóa' hãng 'Toyota' (đang có xe)",
     "Chặn xóa với lỗi 'Không thể xóa hãng đang có xe'",
     "Pass"],
])

# TC16: MANAGE USERS (2)
add_module_heading(doc, "Admin – Quản lý người dùng", "TC16")
add_test_table(doc, HEADERS, [
    ["TC16_01", "Khóa/Mở khóa tài khoản",
     "Click icon 'Khóa' trên 1 user | Xác nhận",
     "User không thể đăng nhập, badge 'Đã khóa' hiển thị",
     "Pass"],
    ["TC16_02", "Xóa người dùng",
     "Click 'Xóa' user không có đơn thuê/mua",
     "User bị xóa, JWT cũ vô hiệu (tăng tokenVersion)",
     "Pass"],
])

# TC17: MANAGE REVIEWS (1)
add_module_heading(doc, "Admin – Quản lý đánh giá", "TC17")
add_test_table(doc, HEADERS, [
    ["TC17_01", "Xem + xóa review vi phạm",
     "Admin /admin/reviews | Click 'Xóa' review spam | Xác nhận",
     "Review bị xóa khỏi DB, không còn hiển thị trên frontend",
     "Pass"],
])

# TC18: MANAGE VIEWING (2)
add_module_heading(doc, "Admin – Quản lý lịch xem xe", "TC18")
add_test_table(doc, HEADERS, [
    ["TC18_01", "Xem tất cả lịch xem xe",
     "/admin/viewing-appointments",
     "Hiển thị tất cả appointment, filter theo status và date range",
     "Pass"],
    ["TC18_02", "Cập nhật trạng thái lịch xem",
     "Chọn appointment | Click 'Đổi trạng thái' | Chọn CONFIRMED",
     "Status cập nhật, gửi email user, slot availability thay đổi",
     "Pass"],
])

# TC19: INVOICE (1)
add_module_heading(doc, "Hóa đơn (Invoices)", "TC19")
add_test_table(doc, HEADERS, [
    ["TC19_01", "Xem chi tiết hóa đơn",
     "Trang /dashboard/rentals | Click icon 'Xem hóa đơn'",
     "Modal hiển thị: mã HĐ, dịch vụ, subtotal, discount, tax 10%, total, status",
     "Pass"],
])

# TC20: SECURITY (4)
add_module_heading(doc, "Bảo mật & Phân quyền", "TC20")
add_test_table(doc, HEADERS, [
    ["TC20_01", "User truy cập admin bị chặn",
     "User (role=user) truy cập /admin",
     "AdminProtectedRoute redirect về /admin/login hoặc 403",
     "Pass"],
    ["TC20_02", "Reset password vô hiệu hóa JWT cũ",
     "Login (token1) → reset password → gọi API bằng token1",
     "HTTP 401 vì tokenVersion không khớp, bắt buộc login lại",
     "Pass"],
    ["TC20_03", "SQL injection prevention",
     "Search: \"'; DROP TABLE users; --\"",
     "JPA parameterized query, không lỗi, hiển thị 'Không tìm thấy'",
     "Pass"],
    ["TC20_04", "XSS prevention",
     "Comment review: '<script>alert(1)</script>'",
     "React escape, hiển thị text thuần, không execute script",
     "Pass"],
])

import os
import sys
os.makedirs(os.path.dirname(OUTPUT_TMP), exist_ok=True)
doc.save(OUTPUT_TMP)
sys.stdout.reconfigure(encoding='utf-8')
print(f"Created: {OUTPUT_TMP}")
print(f"Size: {os.path.getsize(OUTPUT_TMP):,} bytes")
print(f"Please close any Word/WPS that has '{os.path.basename(OUTPUT_PATH)}' open, then rename the new file.")
