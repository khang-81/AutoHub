"""
Tạo bảng kiểm thử AutoHub - 38 test case, GỘP 1 BẢNG DUY NHẤT,
ngôn ngữ đơn giản theo mẫu TC04.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, sys

OUTPUT = r"C:\Do_an\Rent-A-Car\docs\Bang_Kiem_Thu_AutoHub_38cases.docx"

def set_bg(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), color)
    tc_pr.append(shd)

def set_border(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for b in ('top','left','bottom','right'):
        e = OxmlElement(f'w:{b}')
        e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4'); e.set(qn('w:color'),'000000')
        borders.append(e)
    tc_pr.append(borders)

doc = Document()
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = Pt(11)
for s in doc.sections:
    s.top_margin = Cm(1.8); s.bottom_margin = Cm(1.8)
    s.left_margin = Cm(1.5); s.right_margin = Cm(1.5)

t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("BẢNG KIỂM THỬ CHỨC NĂNG"); r.font.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(0x1F,0x3A,0x68)
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Dự án: AutoHub – Hệ thống cho thuê và mua bán xe ô tô")
r.font.bold = True; r.font.size = Pt(13)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Công nghệ: React 19, TypeScript, Tailwind CSS (Frontend) | Spring Boot 3, JPA, JWT (Backend) | SQL Server 2022")
r.font.italic = True; r.font.size = Pt(11)

doc.add_paragraph()

HEADERS = ["ID", "Mô tả", "Dữ liệu đầu vào", "Kỳ vọng", "Kết quả"]

ROWS = [
    # ================== TC01: Đăng nhập / Đăng ký ==================
    ["TC01_01", "Đăng ký tài khoản mới",
     "Email và mật khẩu mới",
     "Tạo được tài khoản và chuyển về trang đăng nhập",
     "Pass"],

    ["TC01_02", "Đăng nhập thành công",
     "Email và mật khẩu đúng",
     "Đăng nhập thành công và vào được trang chủ",
     "Pass"],

    ["TC01_03", "Quên mật khẩu – gửi mã OTP",
     "Email đã đăng ký",
     "Hệ thống gửi mã OTP về email để đặt lại mật khẩu",
     "Pass"],

    ["TC01_04", "Đặt lại mật khẩu bằng OTP",
     "Nhập mã OTP và mật khẩu mới",
     "Đổi mật khẩu thành công, các phiên cũ bị đăng xuất",
     "Pass"],

    ["TC01_05", "Đổi mật khẩu trong dashboard",
     "Vào trang cá nhân, nhập mật khẩu cũ và mới",
     "Đổi mật khẩu thành công và hiện thông báo",
     "Pass"],

    # ================== TC02: Xác minh danh tính ==================
    ["TC02_01", "Upload giấy tờ xác minh (CCCD + GPLX)",
     "Ảnh CCCD và bằng lái xe",
     "Upload thành công, trạng thái chuyển sang đã xác minh",
     "Pass"],

    ["TC02_02", "Chưa xác minh mà thuê xe",
     "User chưa upload giấy tờ, nhấn thuê xe",
     "Hệ thống báo cần xác minh danh tính trước",
     "Pass"],

    ["TC02_03", "Admin duyệt hồ sơ xác minh",
     "Admin vào trang quản lý KYC và nhấn Duyệt",
     "Trạng thái chuyển sang đã duyệt và gửi email cho khách",
     "Pass"],

    # ================== TC03: Danh mục xe ==================
    ["TC03_01", "Tìm kiếm xe theo từ khóa",
     "Nhập tên hãng hoặc tên xe vào ô tìm kiếm",
     "Hiện danh sách các xe phù hợp với từ khóa",
     "Pass"],

    ["TC03_02", "Xem chi tiết một xe",
     "Nhấn vào một xe bất kỳ trong danh sách",
     "Hiện đầy đủ thông tin, hình ảnh và giá của xe",
     "Pass"],

    ["TC03_03", "Lọc xe theo hãng và giá",
     "Chọn hãng Mercedes và khoảng giá từ 500K đến 2 triệu",
     "Chỉ hiện xe Mercedes có giá trong khoảng đã chọn",
     "Pass"],

    # ================== TC04: Đặt lịch xem xe ==================
    ["TC04_01", "Hiển thị danh sách sự kiện",
     "Người dùng đã tạo ít nhất 1 sự kiện",
     "Hệ thống hiển thị danh sách sự kiện theo 4 tab: Sắp diễn ra, Đã qua, Bị từ chối, Chờ duyệt",
     "Pass"],

    ["TC04_02", "Người dùng chưa có sự kiện nào",
     "Người dùng chưa từng tạo sự kiện",
     "Hệ thống hiển thị thông báo \u201cBạn chưa tạo sự kiện nào\u201d",
     "Pass"],

    ["TC04_03", "Lọc sự kiện theo từ khóa",
     "Nhập từ khóa vào ô tìm kiếm (ví dụ: \u201cMercedes\u201d)",
     "Hiển thị danh sách sự kiện phù hợp với từ khóa",
     "Pass"],

    ["TC04_04", "Chuyển tab sự kiện",
     "Chọn tab \u201cĐã qua\u201d",
     "Hệ thống phản hồi trong vòng 2s và hiển thị đúng danh sách",
     "Pass"],

    ["TC04_05", "Đặt lịch xem xe mới",
     "Chọn xe, chọn ngày giờ còn trống và nhấn Đặt lịch",
     "Tạo được lịch xem và gửi email xác nhận",
     "Pass"],

    ["TC04_06", "Hủy lịch xem xe",
     "Mở lịch xem sắp tới và nhấn Hủy",
     "Lịch được hủy và giải phóng khung giờ",
     "Pass"],

    # ================== TC05: Thuê xe ==================
    ["TC05_01", "Chọn ngày thuê xe",
     "Chọn ngày nhận xe và ngày trả xe hợp lệ",
     "Hệ thống tính ra tổng tiền thuê đúng theo số ngày",
     "Pass"],

    ["TC05_02", "Đặt thuê xe thành công",
     "Sau khi chọn ngày và bảo hiểm, nhấn Đặt xe",
     "Tạo được đơn thuê và chuyển sang trang thanh toán",
     "Pass"],

    ["TC05_03", "Hủy đơn thuê xe",
     "Mở đơn thuê đang chờ thanh toán, nhấn Hủy",
     "Đơn được hủy và hoàn cọc theo chính sách",
     "Pass"],

    ["TC05_04", "Trả xe và tính phí phát sinh",
     "Sau ngày trả, nhập số km thực tế và tình trạng nhiên liệu",
     "Hệ thống tính phí trả muộn, phí vượt km, phí thiếu xăng (nếu có)",
     "Pass"],

    # ================== TC06: Thanh toán thuê ==================
    ["TC06_01", "Xác nhận đã chuyển khoản",
     "Trên trang thanh toán, nhấn \u201cTôi đã chuyển khoản\u201d",
     "Đơn chuyển sang trạng thái chờ admin xác nhận",
     "Pass"],

    ["TC06_02", "Admin xác nhận đã nhận tiền",
     "Admin vào đơn thuê và nhấn Xác nhận thanh toán",
     "Đơn thuê được duyệt và gửi email cho khách",
     "Pass"],

    # ================== TC07: Mua xe ==================
    ["TC07_01", "Mua xe thành công",
     "Mở xe đang bán, nhấn Mua ngay",
     "Tạo được đơn mua và chuyển sang trang thanh toán",
     "Pass"],

    ["TC07_02", "Hủy đơn mua xe",
     "Mở đơn mua đang chờ thanh toán, nhấn Hủy",
     "Đơn mua bị hủy và xe được mở bán lại",
     "Pass"],

    ["TC07_03", "Thanh toán đơn mua xe",
     "Trên trang thanh toán đơn mua, nhấn \u201cTôi đã chuyển khoản\u201d",
     "Đơn mua chuyển sang chờ admin xác nhận",
     "Pass"],

    # ================== TC08: Đánh giá ==================
    ["TC08_01", "Viết đánh giá sau khi thuê xe",
     "Chọn số sao và nhập bình luận sau khi hoàn thành đơn thuê",
     "Đánh giá hiển thị trên trang chi tiết xe",
     "Pass"],

    # ================== TC09: Khuyến mãi ==================
    ["TC09_01", "Áp dụng mã giảm giá",
     "Nhập mã khuyến mãi còn hiệu lực vào ô nhập",
     "Được giảm giá đúng theo % của mã",
     "Pass"],

    # ================== TC10: Hồ sơ cá nhân ==================
    ["TC10_01", "Sửa thông tin cá nhân",
     "Thay đổi họ tên, số điện thoại và nhấn Lưu",
     "Thông tin được cập nhật thành công",
     "Pass"],

    # ================== TC11: Liên hệ ==================
    ["TC11_01", "Gửi form liên hệ",
     "Điền họ tên, email, số điện thoại và nội dung",
     "Gửi email cho admin thành công và hiện thông báo",
     "Pass"],

    # ================== TC12: Chatbot AI ==================
    ["TC12_01", "Chat với AI hỏi về xe",
     "Nhấn nút AI và nhập câu hỏi về xe cần tìm",
     "AI trả lời và gợi ý xe phù hợp",
     "Pass"],

    # ================== TC13: Thống kê & Báo cáo ==================
    ["TC13_01", "Xem thống kê tổng quan",
     "Admin vào trang Dashboard",
     "Hiện các số liệu tổng quan và đơn gần đây",
     "Pass"],

    ["TC13_02", "Xuất báo cáo Excel/PDF",
     "Admin nhấn nút Xuất Excel hoặc Xuất PDF",
     "Tải về được file báo cáo đúng định dạng",
     "Pass"],

    # ================== TC14: Admin thêm xe ==================
    ["TC14_01", "Admin thêm xe mới",
     "Điền đầy đủ thông tin xe và nhấn Lưu",
     "Xe mới xuất hiện trong danh sách quản lý xe",
     "Pass"],

    # ================== TC15: Admin thêm hãng ==================
    ["TC15_01", "Admin thêm hãng xe mới",
     "Nhập tên hãng và logo",
     "Hãng mới xuất hiện trong bộ lọc và form thêm xe",
     "Pass"],

    # ================== TC16: Admin quản lý user ==================
    ["TC16_01", "Admin khóa tài khoản người dùng",
     "Nhấn nút Khóa trên một tài khoản user",
     "User đó không đăng nhập lại được",
     "Pass"],

    # ================== TC20: Bảo mật ==================
    ["TC20_01", "User thường vào trang quản trị",
     "User đăng nhập rồi truy cập đường dẫn /admin",
     "Hệ thống chặn và không cho vào trang admin",
     "Pass"],

    ["TC20_02", "Đổi mật khẩu làm hết hạn các phiên cũ",
     "Đăng nhập rồi đổi mật khẩu, dùng lại thiết bị cũ",
     "Thiết bị cũ bị đăng xuất và phải đăng nhập lại",
     "Pass"],
]

table = doc.add_table(rows=1 + len(ROWS), cols=len(HEADERS))
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
col_widths = [Cm(2.0), Cm(5.5), Cm(4.5), Cm(5.5), Cm(1.5)]
for i, w in enumerate(col_widths):
    for c in table.columns[i].cells:
        c.width = w

for j, h in enumerate(HEADERS):
    c = table.rows[0].cells[j]; c.text = h
    set_bg(c, "2E5090"); set_border(c)
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for p in c.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(12)

for i, row in enumerate(ROWS):
    for j, v in enumerate(row):
        c = table.rows[i + 1].cells[j]; c.text = str(v)
        set_border(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(11)

p = doc.add_paragraph()
r = p.add_run("\nGhi chú: Bảng gồm 38 test case cho 17 nhóm chức năng chính. Kết quả \u201cPass\u201d nghĩa là đạt yêu cầu, \u201cFail\u201d là không đạt.")
r.font.italic = True; r.font.size = Pt(10)

os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
doc.save(OUTPUT)
sys.stdout.reconfigure(encoding='utf-8')
print(f"Created: {OUTPUT}")
print(f"Size: {os.path.getsize(OUTPUT):,} bytes")
print(f"Total test cases: {len(ROWS)}")
